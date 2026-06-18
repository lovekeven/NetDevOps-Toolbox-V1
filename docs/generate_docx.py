"""
生成 Word 文档脚本
将 Markdown 设计文档转换为 Word 格式
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_design_document():
    doc = Document()

    # 设置文档样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # 标题样式
    title_style = doc.styles['Title']
    title_style.font.name = '黑体'
    title_style.font.size = Pt(22)
    title_style.font.bold = True

    # 一级标题
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = '黑体'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True

    # 二级标题
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = '黑体'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True

    # 三级标题
    h3_style = doc.styles['Heading 3']
    h3_style.font.name = '黑体'
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True

    # ==================== 封面 ====================
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('SmartNet 智网工坊')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.name = '黑体'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('一站式智能网络运维解决方案')
    run.font.size = Pt(18)
    run.font.name = '黑体'

    doc.add_paragraph('')

    doc_title = doc.add_paragraph()
    doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doc_title.add_run('作品设计文档')
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.name = '黑体'

    doc.add_paragraph('')
    doc.add_paragraph('')

    # 文档信息
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('所在赛道与赛项：A')
    run.font.size = Pt(14)

    info2 = doc.add_paragraph()
    info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info2.add_run('文档版本：V1.0')
    run.font.size = Pt(12)

    info3 = doc.add_paragraph()
    info3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info3.add_run('编写日期：2026年6月17日')
    run.font.size = Pt(12)

    doc.add_page_break()

    # ==================== 目录 ====================
    doc.add_heading('目录', level=1)

    toc_items = [
        '一、目标问题与意义价值',
        '二、设计思路与方案',
        '三、方案实现',
        '四、运行结果/应用效果',
        '五、创新与特色',
        '附录'
    ]

    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(item)
        run.font.size = Pt(14)

    doc.add_page_break()

    # ==================== 一、目标问题与意义价值 ====================
    doc.add_heading('一、目标问题与意义价值', level=1)

    doc.add_heading('1.1 应用领域', level=2)
    p = doc.add_paragraph()
    p.add_run('本作品面向企业网络运维领域，为网络运维工程师、系统管理员、IT运维团队提供一站式智能网络运维解决方案。')

    doc.add_heading('1.2 解决的问题', level=2)
    p = doc.add_paragraph()
    p.add_run('核心问题：')

    problems = [
        '拓扑发现效率低：传统人工绘制网络拓扑耗时耗力，且容易出错',
        '设备管理分散：多厂商设备需要不同工具管理，运维复杂度高',
        '故障排查慢：缺乏智能化诊断手段，故障定位依赖人工经验',
        '运维成本高：需要安装多种客户端软件，学习成本高'
    ]

    for problem in problems:
        p = doc.add_paragraph()
        p.style = doc.styles['List Bullet']
        p.add_run(problem)

    doc.add_heading('1.3 实现目标', level=2)

    goals = [
        '自动化拓扑发现：通过SNMP/LLDP协议自动发现网络设备，生成可视化拓扑图',
        '统一设备管理：支持多厂商设备统一管理，提供Web SSH终端',
        '智能监控告警：实时监控设备状态，异常自动告警',
        'AI辅助运维：利用人工智能技术辅助故障诊断和运维决策'
    ]

    for goal in goals:
        p = doc.add_paragraph()
        p.style = doc.styles['List Bullet']
        p.add_run(goal)

    doc.add_heading('1.4 基本功能', level=2)

    # 功能表格
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'

    headers = ['功能模块', '功能说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    functions = [
        ('拓扑管理', '自动发现网络设备，生成可视化拓扑图'),
        ('运维工具', 'Web SSH终端、一键设备查询、常用命令集'),
        ('监控告警', '实时监控设备状态，灵活告警规则配置'),
        ('AI分析', '设备健康分析、拓扑推测、报告生成'),
        ('设备管理', '设备增删改查、健康检查、配置备份')
    ]

    for i, (module, desc) in enumerate(functions):
        table.rows[i+1].cells[0].text = module
        table.rows[i+1].cells[1].text = desc

    doc.add_heading('1.5 理论意义与应用价值', level=2)
    p = doc.add_paragraph()
    p.add_run('理论意义：')

    meanings = [
        '探索AI技术在网络运维领域的应用模式',
        '研究多厂商设备统一管理的技术方案',
        '实现Web终端与网络设备的实时交互'
    ]

    for meaning in meanings:
        p = doc.add_paragraph()
        p.style = doc.styles['List Bullet']
        p.add_run(meaning)

    p = doc.add_paragraph()
    p.add_run('应用价值：')

    values = [
        '提高网络运维效率，降低人工成本',
        '降低运维门槛，新手也能快速上手',
        '适应企业多厂商设备环境，降低管理复杂度',
        '为智能化运维提供技术参考'
    ]

    for value in values:
        p = doc.add_paragraph()
        p.style = doc.styles['List Bullet']
        p.add_run(value)

    doc.add_page_break()

    # ==================== 二、设计思路与方案 ====================
    doc.add_heading('二、设计思路与方案', level=1)

    doc.add_heading('2.1 设计思路', level=2)
    p = doc.add_paragraph()
    p.add_run('核心理念：以用户体验为中心，将复杂的网络运维操作简化为可视化、智能化的交互流程。')

    p = doc.add_paragraph()
    p.add_run('设计原则：')

    principles = [
        '简洁易用：界面简洁，操作直观，降低学习成本',
        '功能全面：覆盖网络运维的核心场景',
        '智能化：引入AI技术，提供智能决策支持',
        '可扩展：支持多厂商设备，便于功能扩展'
    ]

    for principle in principles:
        p = doc.add_paragraph()
        p.style = doc.styles['List Bullet']
        p.add_run(principle)

    doc.add_heading('2.2 技术路线', level=2)
    p = doc.add_paragraph()
    p.add_run('本作品采用前后端分离架构，前端使用HTML5+CSS3+JavaScript，后端使用Python Flask框架，通过RESTful API和WebSocket实现数据交互。')

    doc.add_heading('2.3 详细设计方案', level=2)

    doc.add_heading('2.3.1 拓扑发现方案', level=3)
    p = doc.add_paragraph()
    p.add_run('技术方案：')

    solutions = [
        'SNMP协议采集：通过pysnmp库实现SNMP协议，获取设备信息',
        'LLDP邻居发现：读取设备LLDP邻居表，获取连接关系',
        '多层递归扫描：从种子设备开始，递归扫描全网设备',
        'vis.js可视化：使用vis.js库渲染拓扑图'
    ]

    for solution in solutions:
        p = doc.add_paragraph()
        p.style = doc.styles['List Bullet']
        p.add_run(solution)

    p = doc.add_paragraph()
    p.add_run('辅助功能：MAC地址表分析')
    p = doc.add_paragraph()
    p.add_run('当LLDP协议不可用时（如某些模拟器环境），可通过分析设备的MAC地址表来推导连接关系。此功能作为拓扑发现的辅助手段，适用于特殊环境。')

    doc.add_heading('2.3.2 运维工具方案', level=3)
    p = doc.add_paragraph()
    p.add_run('Web SSH终端：')

    solutions = [
        '前端：xterm.js终端模拟器',
        '后端：paramiko SSH连接库',
        '通信：WebSocket实时数据传输'
    ]

    for solution in solutions:
        p = doc.add_paragraph()
        p.style = doc.styles['List Bullet']
        p.add_run(solution)

    doc.add_heading('2.3.3 监控告警方案', level=3)
    p = doc.add_paragraph()
    p.add_run('监控方式：')

    methods = [
        '定时任务：APScheduler定时执行',
        '数据采集：SNMP/SSH获取设备状态',
        '实时推送：WebSocket推送监控数据'
    ]

    for method in methods:
        p = doc.add_paragraph()
        p.style = doc.styles['List Bullet']
        p.add_run(method)

    doc.add_heading('2.3.4 AI分析方案', level=3)
    p = doc.add_paragraph()
    p.add_run('技术方案：')

    solutions = [
        'AI模型：DeepSeek大语言模型',
        '调用方式：API调用',
        '应用场景：健康分析、拓扑推测、报告生成'
    ]

    for solution in solutions:
        p = doc.add_paragraph()
        p.style = doc.styles['List Bullet']
        p.add_run(solution)

    doc.add_page_break()

    # ==================== 三、方案实现 ====================
    doc.add_heading('三、方案实现', level=1)

    doc.add_heading('3.1 技术实现', level=2)

    doc.add_heading('3.1.1 前端技术', level=3)

    # 前端技术表格
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'

    headers = ['技术', '用途']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    tech = [
        ('HTML5 + CSS3', '页面结构和样式'),
        ('Bootstrap 5', 'UI组件库'),
        ('vis.js', '拓扑图可视化'),
        ('ECharts', '数据图表'),
        ('xterm.js', 'Web终端')
    ]

    for i, (name, usage) in enumerate(tech):
        table.rows[i+1].cells[0].text = name
        table.rows[i+1].cells[1].text = usage

    doc.add_heading('3.1.2 后端技术', level=3)

    # 后端技术表格
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'

    headers = ['技术', '用途']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    tech = [
        ('Flask', 'Web框架'),
        ('Flask-SocketIO', 'WebSocket支持'),
        ('pysnmp', 'SNMP协议'),
        ('paramiko', 'SSH连接'),
        ('Nornir', '并发框架'),
        ('SQLite', '数据库'),
        ('DeepSeek API', 'AI分析')
    ]

    for i, (name, usage) in enumerate(tech):
        table.rows[i+1].cells[0].text = name
        table.rows[i+1].cells[1].text = usage

    doc.add_heading('3.2 功能实现', level=2)

    doc.add_heading('3.2.1 拓扑管理功能', level=3)
    p = doc.add_paragraph()
    p.add_run('通过SNMP/LLDP协议自动发现网络设备，生成可视化拓扑图。支持拓扑编辑、自动布局、多格式导出等功能。')

    doc.add_heading('3.2.2 运维工具功能', level=3)
    p = doc.add_paragraph()
    p.add_run('提供Web SSH终端、一键设备查询、常用命令集等功能。支持多会话管理、快速连接设备、命令模板支持。')

    doc.add_heading('3.2.3 监控告警功能', level=3)
    p = doc.add_paragraph()
    p.add_run('实时监控设备状态，支持CPU、内存、接口等指标监控。灵活的告警规则配置，支持邮件告警通知。')

    doc.add_heading('3.2.4 AI分析功能', level=3)
    p = doc.add_paragraph()
    p.add_run('利用AI技术分析设备健康数据，给出诊断建议。支持AI拓扑推测、自动报告生成等功能。')

    doc.add_page_break()

    # ==================== 四、运行结果/应用效果 ====================
    doc.add_heading('四、运行结果/应用效果', level=1)

    doc.add_heading('4.1 系统运行环境', level=2)

    # 运行环境表格
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    headers = ['项目', '配置']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    env = [
        ('操作系统', 'Windows 11'),
        ('Python版本', '3.10+'),
        ('浏览器', 'Chrome/Edge/Firefox'),
        ('网络环境', '局域网/互联网')
    ]

    for i, (item, config) in enumerate(env):
        table.rows[i+1].cells[0].text = item
        table.rows[i+1].cells[1].text = config

    doc.add_heading('4.2 功能测试结果', level=2)

    doc.add_heading('4.2.1 拓扑发现测试', level=3)
    p = doc.add_paragraph()
    p.add_run('测试场景：模拟器环境（EVE-NG）')

    # 测试结果表格
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['测试项', '结果', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    results = [
        ('设备发现', '成功', '能够发现所有配置SNMP的设备'),
        ('链路发现', '成功', '能够通过LLDP发现设备连接关系'),
        ('拓扑渲染', '成功', 'vis.js正确渲染拓扑图'),
        ('多厂商支持', '成功', '支持H3C、Huawei设备')
    ]

    for i, (test, result, desc) in enumerate(results):
        table.rows[i+1].cells[0].text = test
        table.rows[i+1].cells[1].text = result
        table.rows[i+1].cells[2].text = desc

    doc.add_heading('4.2.2 运维工具测试', level=3)
    p = doc.add_paragraph()
    p.add_run('测试场景：SSH连接网络设备')

    # 测试结果表格
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['测试项', '结果', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    results = [
        ('SSH连接', '成功', '能够连接H3C、Huawei设备'),
        ('命令执行', '成功', '能够执行命令并返回结果'),
        ('多会话', '成功', '支持同时连接多台设备'),
        ('快速连接', '成功', '可以从设备列表快速连接')
    ]

    for i, (test, result, desc) in enumerate(results):
        table.rows[i+1].cells[0].text = test
        table.rows[i+1].cells[1].text = result
        table.rows[i+1].cells[2].text = desc

    doc.add_heading('4.3 性能测试结果', level=2)

    # 性能测试表格
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    headers = ['指标', '结果']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    perf = [
        ('页面加载时间', '< 2秒'),
        ('拓扑扫描时间', '10-30秒（取决于设备数量）'),
        ('SSH连接时间', '< 1秒'),
        ('并发设备数', '支持10+台设备同时管理')
    ]

    for i, (metric, result) in enumerate(perf):
        table.rows[i+1].cells[0].text = metric
        table.rows[i+1].cells[1].text = result

    doc.add_heading('4.4 应用效果', level=2)

    effects = [
        '提高运维效率：拓扑自动发现替代手工绘制，效率提升80%以上',
        '降低运维门槛：Web终端无需安装客户端，新手也能快速上手',
        '统一管理界面：多厂商设备统一管理，降低学习成本',
        '智能化决策：AI辅助分析，提高故障诊断效率'
    ]

    for effect in effects:
        p = doc.add_paragraph()
        p.style = doc.styles['List Bullet']
        p.add_run(effect)

    doc.add_page_break()

    # ==================== 五、创新与特色 ====================
    doc.add_heading('五、创新与特色', level=1)

    innovations = [
        {
            'title': 'AI智能运维决策系统',
            'desc': '将AI技术深度集成到网络运维中，实现智能化决策支持。包括AI健康分析、AI拓扑推测、AI报告生成、AI故障诊断等功能。',
            'advantage': '提高运维效率，降低人工成本，智能化决策支持'
        },
        {
            'title': 'Web终端零客户端管理',
            'desc': '将SSH终端集成到Web界面，实现零客户端管理。基于xterm.js的全功能终端模拟，支持多会话管理、快速连接设备、命令模板支持。',
            'advantage': '无需安装客户端，多设备并行管理，提高运维效率'
        },
        {
            'title': '多厂商设备统一管理',
            'desc': '提供统一的设备管理接口，支持H3C、Huawei、Cisco等多厂商设备。统一的设备管理接口、多厂商命令模板库、统一的健康检查标准。',
            'advantage': '降低运维复杂度，提高管理效率，适应企业多厂商环境'
        },
        {
            'title': '实时监控与智能告警',
            'desc': '实时监控设备状态，智能告警通知。WebSocket实时推送监控数据，灵活的告警规则配置，多种告警级别支持，邮件告警通知。',
            'advantage': '实时掌握设备状态，快速发现和处理故障，灵活的告警策略'
        },
        {
            'title': '可视化拓扑管理',
            'desc': '自动发现网络设备，生成可视化拓扑图。SNMP/LLDP协议自动发现，vis.js可视化渲染，拓扑编辑功能，多格式导出支持。',
            'advantage': '自动化拓扑发现，替代手工绘制，可视化展示，直观清晰'
        }
    ]

    for i, innovation in enumerate(innovations, 1):
        doc.add_heading(f'5.{i} {innovation["title"]}', level=2)
        p = doc.add_paragraph()
        p.add_run('创新点：')
        p = doc.add_paragraph()
        p.add_run(innovation['desc'])
        p = doc.add_paragraph()
        p.add_run('技术优势：')
        p = doc.add_paragraph()
        p.add_run(innovation['advantage'])

    doc.add_page_break()

    # ==================== 附录 ====================
    doc.add_heading('附录', level=1)

    doc.add_heading('A. 项目地址', level=2)
    p = doc.add_paragraph()
    p.add_run('GitHub仓库：https://github.com/lovekeven/NetDevOps-Toolbox-V1')

    doc.add_heading('B. 开发环境', level=2)

    # 开发环境表格
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    headers = ['项目', '配置']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    env = [
        ('操作系统', 'Windows 11'),
        ('Python版本', '3.10+'),
        ('开发工具', 'VS Code'),
        ('版本控制', 'Git')
    ]

    for i, (item, config) in enumerate(env):
        table.rows[i+1].cells[0].text = item
        table.rows[i+1].cells[1].text = config

    doc.add_heading('C. 参考文献', level=2)

    refs = [
        '[1] RFC 1213 - MIB-II',
        '[2] RFC 2011 - SNMPv2 MIB',
        '[3] IEEE 802.1AB - LLDP',
        '[4] Flask官方文档 - https://flask.palletsprojects.com/',
        '[5] vis.js官方文档 - https://visjs.org/'
    ]

    for ref in refs:
        p = doc.add_paragraph()
        p.add_run(ref)

    # 文档信息
    doc.add_paragraph('')
    doc.add_paragraph('')

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    info = [
        ('文档版本', 'V1.0'),
        ('编写日期', '2026年6月17日'),
        ('所在赛道', 'A赛道'),
        ('作品名称', 'SmartNet 智网工坊')
    ]

    for i, (item, value) in enumerate(info):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = value

    # 保存文档
    doc.save('docs/作品设计文档.docx')
    print('Word 文档已生成：docs/作品设计文档.docx')

if __name__ == '__main__':
    create_design_document()
